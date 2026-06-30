"""Utility per la generazione di documenti per il Menu Creation Studio."""

from collections import defaultdict
from io import BytesIO
from typing import Tuple
import zipfile

from django.template.loader import render_to_string
from django.utils.text import slugify

import weasyprint

from .models import Allergene

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:  # pragma: no cover - gestito a runtime
    Document = None
    Pt = None


def _ensure_python_docx():
    if Document is None:
        raise RuntimeError(
            "python-docx non è installato. Aggiungere 'python-docx' ai requirements per generare documenti Word."
        )


def _load_document_from_template(template_field):
    _ensure_python_docx()
    if template_field and getattr(template_field, 'name', None):
        return Document(template_field.path)
    return Document()


def _format_allergen_list(piatto):
    allergeni = [a.nome for a in piatto.allergeni.all()]
    if piatto.allergen_summary:
        allergeni.append(piatto.allergen_summary)
    return ", ".join(allergeni) if allergeni else "Nessuno"


def _group_piatti(menu):
    grouped = defaultdict(list)
    qs = menu.piatti.all().select_related('base_item').prefetch_related('ingredienti', 'allergeni')
    for piatto in qs:
        grouped[piatto.get_categoria_display()].append(piatto)
    return grouped


def build_menu_docx(menu) -> Tuple[BytesIO, str]:
    """Genera un documento Word con il menu completo rispettando il layout."""

    document = _load_document_from_template(getattr(menu.layout, 'documento_word', None))
    layout = menu.layout
    if not layout:
        # Fallback se non c'è layout
        document.add_heading(menu.nome, level=1)
        grouped = _group_piatti(menu)
        for categoria, piatti in grouped.items():
            document.add_heading(categoria, level=2)
            for piatto in piatti:
                document.add_paragraph(piatto.nome, style='List Bullet' if 'List Bullet' in document.styles else None)
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer, f"{slugify(menu.nome)}_menu.docx"

    # Layout dinamico
    blocks_config = layout.struttura_blocchi.get('blocks', {})
    order = layout.struttura_blocchi.get('order', {'1': ['logo', 'info', 'sections']})

    # In Word le colonne sono gestite in modo sequenziale per ora
    all_blocks = list(order.get('1', [])) + list(order.get('2', []))

    for block_id in all_blocks:
        if not blocks_config.get(block_id, {}).get('enabled', True):
            continue

        if block_id == 'logo':
            # Il logo di solito è nell'header del template, saltiamo l'aggiunta inline
            pass
        elif block_id == 'info':
            document.add_heading(menu.nome, level=1)
            if menu.data_evento:
                document.add_paragraph(f"Data: {menu.data_evento:%d/%m/%Y}")
            document.add_paragraph(f"Turno: {menu.get_turno_display() if hasattr(menu, 'get_turno_display') else menu.turno}")
            if menu.ospiti_target:
                document.add_paragraph(f"Ospiti: {menu.ospiti_target}")
        elif block_id == 'sections':
            grouped = _group_piatti(menu)
            for categoria, piatti in grouped.items():
                document.add_heading(categoria, level=2)
                for piatto in piatti:
                    p = document.add_paragraph()
                    p.style = document.styles['List Bullet'] if 'List Bullet' in document.styles else p.style
                    run = p.add_run(piatto.nome)
                    run.bold = True
                    if piatto.prezzo:
                        p.add_run(f" \t {piatto.prezzo} €") # Tab per allineamento a destra (dipende dai tab stop del template)

                    if piatto.descrizione:
                        document.add_paragraph(piatto.descrizione)
                    allergeni = _format_allergen_list(piatto)
                    document.add_paragraph(f"Allergeni: {allergeni}", style='Caption' if 'Caption' in document.styles else None)
        elif block_id == 'legend':
            allergeni = Allergene.objects.filter(piatti__menu=menu).distinct()
            if allergeni.exists():
                document.add_heading("Legenda Allergeni", level=2)
                for allergene in allergeni:
                    p = document.add_paragraph(style='Caption' if 'Caption' in document.styles else None)
                    p.add_run(f"{allergene.nome}: ").bold = True
                    if allergene.descrizione:
                        p.add_run(allergene.descrizione)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    filename = f"{slugify(menu.nome)}_menu.docx"
    return buffer, filename


def build_cavalieri_docx(menu) -> Tuple[BytesIO, str]:
    """Genera un documento Word con i cavalieri per ogni piatto rispettando la configurazione."""

    template_field = getattr(menu.cavaliere_template, 'documento_word', None) if menu.cavaliere_template else None
    config = menu.cavaliere_template.configurazione if menu.cavaliere_template else {}
    document = _load_document_from_template(template_field)

    piatti = list(menu.piatti.all().prefetch_related('allergeni'))
    for index, piatto in enumerate(piatti):
        # Nome Piatto
        h = document.add_heading(piatto.nome, level=1)
        if Pt and 'font_size_nome' in config:
            for run in h.runs:
                run.font.size = Pt(config['font_size_nome'])

        # Descrizione (se abilitata)
        if config.get('mostra_descrizione', True) and piatto.descrizione:
            p_desc = document.add_paragraph(piatto.descrizione)
            if Pt and 'font_size_desc' in config:
                for run in p_desc.runs:
                    run.font.size = Pt(config['font_size_desc'])

        # Prezzo (se abilitato)
        if config.get('mostra_prezzo', False) and piatto.prezzo:
            p_prezzo = document.add_paragraph(f"Prezzo: {piatto.prezzo} €")
            if Pt and 'font_size_prezzo' in config:
                for run in p_prezzo.runs:
                    run.font.size = Pt(config['font_size_prezzo'])

        # Allergeni (se abilitati)
        if config.get('mostra_allergeni', True):
            allergeni = _format_allergen_list(piatto)
            p_all = document.add_paragraph(f"Allergeni: {allergeni}")
            if Pt:
                size = config.get('font_size_allergeni', 10)
                for run in p_all.runs:
                    run.font.size = Pt(size)

        # Immagine piatto (se disponibile e abilitata)
        if config.get('mostra_immagine', False) and piatto.immagine:
            try:
                document.add_picture(piatto.immagine.path, width=Pt(150))
            except Exception:
                pass

        if index < len(piatti) - 1:
            document.add_page_break()

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    filename = f"{slugify(menu.nome)}_cavalieri.docx"
    return buffer, filename


def build_menu_pdf(menu, doc_type="menu") -> Tuple[bytes, str]:
    """Renderizza un PDF con WeasyPrint per menu o cavalieri."""

    template_path = (
        "menu_generator/pdf/cavaliere_template.html"
        if doc_type == "cavaliere"
        else "menu_generator/pdf/menu_template.html"
    )

    # Fallback per layout mancante per evitare errori nei template
    layout = menu.layout
    if not layout:
        layout = {
            'struttura_blocchi': {
                'blocks': {
                    'logo': {'enabled': True},
                    'info': {'enabled': True},
                    'sections': {'enabled': True},
                    'legend': {'enabled': True},
                },
                'order': {'1': ['logo', 'info', 'sections', 'legend'], '2': []},
                'columns': 1
            },
            'font_principale': 'serif',
            'colore_font': '#1a1a1a'
        }

    context = {
        "menu": menu,
        "layout": layout,
        "piatti_per_categoria": _group_piatti(menu),
        "allergeni": Allergene.objects.filter(piatti__menu=menu).distinct(),
    }
    html_string = render_to_string(template_path, context)
    pdf_bytes = weasyprint.HTML(string=html_string).write_pdf()
    filename = f"{slugify(menu.nome)}_{doc_type}.pdf"
    return pdf_bytes, filename


def build_menu_bundle(menu, doc_type="menu", include_cavalieri=False) -> Tuple[BytesIO, str]:
    """Costruisce un archivio zip con PDF e documenti Word."""

    pdf_bytes, pdf_filename = build_menu_pdf(menu, doc_type)
    bundle = BytesIO()
    with zipfile.ZipFile(bundle, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(pdf_filename, pdf_bytes)
        if doc_type == "cavaliere" or include_cavalieri:
            cav_buffer, cav_filename = build_cavalieri_docx(menu)
            archive.writestr(cav_filename, cav_buffer.getvalue())
        menu_buffer, menu_filename = build_menu_docx(menu)
        archive.writestr(menu_filename, menu_buffer.getvalue())

    bundle.seek(0)
    filename = f"{slugify(menu.nome)}_{doc_type}_bundle.zip"
    return bundle, filename
